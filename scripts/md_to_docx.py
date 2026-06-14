"""
Minimal, dependency-light Markdown -> .docx converter (python-docx) for the repo's reports.
Handles: #/##/### headings, **bold** inline, pipe tables (with header), - bullets, 1. numbers,
--- rules, and paragraphs. Not a full CommonMark engine — just enough for our docs.

Run: .venv/bin/python scripts/md_to_docx.py docs/EDGE_CASE_STUDIES.md [out.docx]
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _add_runs(paragraph, text: str):
    """Split on **bold** and add runs, preserving inline bold."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = (i % 2 == 1)  # odd chunks were inside ** **


def _emit_table(doc, rows):
    """rows: list of cell-lists; row 0 is the header (separator row already removed)."""
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, cells in enumerate(rows):
        tr = t.add_row().cells
        for ci in range(ncol):
            txt = cells[ci] if ci < len(cells) else ""
            p = tr[ci].paragraphs[0]
            _add_runs(p, txt.strip())
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def _split_row(line: str):
    parts = line.strip().strip("|").split("|")
    return [c.strip() for c in parts]


def convert(md_path: Path, out_path: Path):
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    lines = md_path.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # table block
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[i + 1]):
            block = [_split_row(line)]
            i += 2  # skip header + separator
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(_split_row(lines[i]))
                i += 1
            _emit_table(doc, block)
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "---":
            doc.add_paragraph("")  # light separator
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)
        i += 1
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[done] wrote {out_path} ({out_path.stat().st_size:,} bytes)")


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: md_to_docx.py <input.md> [output.docx]   OR   md_to_docx.py a.md b.md c.md",
              file=sys.stderr)
        return 1
    args = sys.argv[1:]
    # Two-arg explicit form ONLY when the 2nd path is a .docx target; otherwise treat every
    # argument as an input .md and write each to its own .docx (prevents clobbering an .md by
    # passing it as argv[2] — the bug that overwrote STRATEGY_TEST_PLAN.md with binary).
    if len(args) == 2 and args[1].lower().endswith(".docx"):
        convert(Path(args[0]), Path(args[1]))
        return 0
    for a in args:
        src = Path(a)
        if src.suffix.lower() != ".md":
            print(f"[skip] {a} (not a .md input)", file=sys.stderr)
            continue
        convert(src, src.with_suffix(".docx"))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
